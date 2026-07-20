from pathlib import Path
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ============================================================
# Paths
# ============================================================

ROOT_DIR = Path(
    r"\outputs\data_figures"
)

OUTPUT_DOC = Path(
    r"\outputs\fNIRS_Assignment_Report.docx"
)

# ============================================================
# Find Subject Folders
# ============================================================

subject_folders = sorted(
    (
        folder
        for folder in ROOT_DIR.iterdir()
        if folder.is_dir()
        and folder.name.startswith("Subject-")
    ),
    key=lambda folder: int(folder.name.split("-")[1]),
)

print("=" * 70)
print("Generating fNIRS Assignment Report")
print("=" * 70)

print(f"Input directory : {ROOT_DIR}")
print(f"Output report   : {OUTPUT_DOC}")
print(f"Subjects found  : {len(subject_folders)}")

if len(subject_folders) == 0:
    raise FileNotFoundError(
        f"No subject folders were found in:\n{ROOT_DIR}"
    )
    
# ============================================================
# Create Word Document
# ============================================================

document = Document()

# ------------------------------------------------------------
# Page Layout
# ------------------------------------------------------------

section = document.sections[0]

section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)

section.left_margin = Inches(0.6)
section.right_margin = Inches(0.6)

# ------------------------------------------------------------
# Title
# ------------------------------------------------------------

title = document.add_heading(
    "fNIRS Data Processing Assignment",
    level=0,
)

title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ------------------------------------------------------------
# Candidate Details
# ------------------------------------------------------------

info = document.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER

run = info.add_run("Candidate: Ritika Jain\n")
run.bold = True
run.font.size = Pt(12)

info.add_run(
    "BS-MS (Data Science and Engineering)\n"
)

info.add_run(
    "Indian Institute of Science Education and Research (IISER) Bhopal\n"
)

info.add_run(
    "Date: July 2026\n\n"
)

subtitle = info.add_run(
    "Automated fNIRS Preprocessing and Visualization Report"
)

subtitle.italic = True
subtitle.font.size = Pt(11)

# ------------------------------------------------------------
# Separator
# ------------------------------------------------------------

document.add_paragraph("=" * 70)

# ------------------------------------------------------------
# Objective
# ------------------------------------------------------------

document.add_heading(
    "Objective",
    level=1,
)

document.add_paragraph(
    "The objective of this assignment is to preprocess the provided "
    "fNIRS recordings, extract task-related haemodynamic responses, "
    "and visualize cortical activation patterns for all annotated "
    "experimental conditions. An automated MNE-Python pipeline was "
    "developed to process every subject consistently and generate "
    "Before (-2–0 s), During (0–5 s), and After (5–10 s) activation "
    "maps for both oxygenated (HbO) and deoxygenated (HbR) haemoglobin."
)

# ------------------------------------------------------------
# Processing Pipeline
# ------------------------------------------------------------

document.add_heading(
    "Processing Pipeline",
    level=1,
)

pipeline = (
    "Raw SNIRF Recording\n"
    "        ↓\n"
    "Optical Density\n"
    "        ↓\n"
    "Beer-Lambert Law\n"
    "        ↓\n"
    "TDDR Motion Correction\n"
    "        ↓\n"
    "Band-pass Filtering (0.01–0.10 Hz)\n"
    "        ↓\n"
    "Event Extraction\n"
    "        ↓\n"
    "Epoching (-2 s to 15 s)\n"
    "        ↓\n"
    "Evoked Response Computation\n"
    "        ↓\n"
    "HbO / HbR Separation\n"
    "        ↓\n"
    "Before / During / After Activation Maps"
)

paragraph = document.add_paragraph()

run = paragraph.add_run(pipeline)
run.font.name = "Consolas"
run.font.size = Pt(10)

# ------------------------------------------------------------
# Start Results on New Page
# ------------------------------------------------------------

document.add_page_break()

# ============================================================
# Helper Functions
# ============================================================

def add_heading(text, level=1):
    """
    Add a heading to the document.
    """
    document.add_heading(text, level=level)
        
def read_csv(csv_file):
    """
    Read a CSV file if it exists.
    """

    csv_file = Path(csv_file)

    if csv_file.exists():
        return pd.read_csv(csv_file)

    return pd.DataFrame()

def add_dataframe(df):
    """
    Insert a pandas DataFrame as a formatted Word table.
    """

    if df.empty:
        document.add_paragraph("No data available.")
        return

    table = document.add_table(
        rows=1,
        cols=len(df.columns),
        style="Table Grid",
    )

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # --------------------------------------------------------
    # Header
    # --------------------------------------------------------

    header_cells = table.rows[0].cells

    for i, column in enumerate(df.columns):

        paragraph = header_cells[i].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(str(column))
        run.bold = True

    # --------------------------------------------------------
    # Data Rows
    # --------------------------------------------------------

    for _, row in df.iterrows():

        cells = table.add_row().cells

        for j, value in enumerate(row):

            paragraph = cells[j].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if isinstance(value, float):

                # Scientific notation for very small values
                if abs(value) < 1e-4:
                    text = f"{value:.6e}"
                else:
                    text = f"{value:.6f}"

            else:
                text = str(value)

            paragraph.add_run(text)

    document.add_paragraph()
            
def add_activation_images(condition_folder, signal):
    """
    Insert Before, During and After activation maps
    into a three-column table.
    """

    document.add_heading(signal, level=3)

    table = document.add_table(
        rows=2,
        cols=3,
        style="Table Grid",
    )

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "Before",
        "During",
        "After",
    ]

    # --------------------------------------------------------
    # Header Row
    # --------------------------------------------------------

    for i, text in enumerate(headers):

        cell = table.cell(0, i)

        paragraph = cell.paragraphs[0]

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        run = paragraph.add_run(text)

        run.bold = True

    # --------------------------------------------------------
    # Images
    # --------------------------------------------------------

    image_names = [
        f"{signal}_Before.png",
        f"{signal}_During.png",
        f"{signal}_After.png",
    ]

    for i, image_name in enumerate(image_names):

        image = condition_folder / image_name

        paragraph = table.cell(1, i).paragraphs[0]

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        if image.exists():

            paragraph.add_run().add_picture(
                str(image),
                width=Inches(1.9),
            )

        else:

            paragraph.add_run(
                "Image not found"
            )

    document.add_paragraph()

    caption = document.add_paragraph()

    caption.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    caption_run = caption.add_run(
        f"{signal} activation maps "
        "(Before, During and After)"
    )

    caption_run.italic = True

    document.add_paragraph()
            
    
def add_summary_tables(subject_folder):
    """
    Insert trial and activation summary tables.
    """

    trial_csv = (
        subject_folder /
        "trial_summary.csv"
    )

    activation_csv = (
        subject_folder /
        "activation_summary.csv"
    )

    document.add_heading(
        "Trial Summary",
        level=2,
    )

    add_dataframe(
        read_csv(trial_csv)
    )

    document.add_heading(
        "Activation Summary",
        level=2,
    )

    add_dataframe(
        read_csv(activation_csv)
    )

# ============================================================
# Generate Report
# ============================================================

for subject_folder in subject_folders:

    subject_id = subject_folder.name.split("-")[1]

    print(f"Processing Subject {subject_id}")

    # ========================================================
    # Subject Heading
    # ========================================================

    subject_heading = document.add_heading(
        f"Subject {subject_id}",
        level=1,
    )

    subject_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = subject_heading.runs[0]
    run.bold = True
    run.font.size = Pt(18)

    document.add_paragraph(
        "Haemodynamic activation maps for all annotated "
        "experimental conditions."
    )

    # ========================================================
    # Process Conditions
    # ========================================================

    for cond in ["1", "2", "4", "6"]:

        condition_folder = (
            subject_folder /
            f"Condition_{cond}"
        )

        if not condition_folder.exists():
            continue

        document.add_heading(
            f"Condition {cond}",
            level=2,
        )

        document.add_paragraph(
            "Average cortical activation before, during "
            "and after the annotated task event."
        )

        # ----------------------------------------------------
        # HbO
        # ----------------------------------------------------

        add_activation_images(
            condition_folder,
            "HbO",
        )

        # ----------------------------------------------------
        # HbR
        # ----------------------------------------------------

        add_activation_images(
            condition_folder,
            "HbR",
        )

        # ----------------------------------------------------
        # Separator
        # ----------------------------------------------------

        document.add_paragraph("=" * 80)

    # ========================================================
    # Subject Summary
    # ========================================================

    add_summary_tables(
        subject_folder
    )
    
# ============================================================
# Final Remarks
# ============================================================

document.add_heading(
    "Report Summary",
    level=1,
)

document.add_paragraph(
    "This report presents the preprocessing results for all available "
    "subjects in the provided fNIRS emotion recognition dataset. "
    "For each subject and each annotated experimental condition "
    "(1, 2, 4 and 6), Before, During and After activation maps are "
    "provided for both HbO and HbR signals together with summary "
    "tables of the extracted haemodynamic responses."
)

# ============================================================
# Save Report
# ============================================================

document.save(OUTPUT_DOC)

print("\n" + "=" * 70)
print("Report Generated Successfully")
print("=" * 70)
print(f"Report saved to : {OUTPUT_DOC}")
